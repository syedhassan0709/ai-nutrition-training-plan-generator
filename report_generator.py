"""
Report Generator Module for Nutrition App
Populates Word templates with AI-generated content and charts.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os
import logging
from typing import Dict, Any, Optional
from datetime import datetime
import re


class ReportGenerator:
    """Populates Word templates with AI-generated content."""
    
    def __init__(self, templates_dir: str = "templates", output_dir: str = "output"):
        """
        Initialize ReportGenerator.
        
        Args:
            templates_dir (str): Directory containing Word templates
            output_dir (str): Directory to save generated reports
        """
        self.templates_dir = templates_dir
        self.output_dir = output_dir
        self.logger = logging.getLogger(__name__)
        
        # Create directories if they don't exist
        os.makedirs(templates_dir, exist_ok=True)
        os.makedirs(output_dir, exist_ok=True)
    
    def generate_all_reports(self, 
                           questionnaire_data: Dict[str, Any],
                           llm_content: Dict[str, str],
                           chart_path: Optional[str] = None) -> Dict[str, str]:
        """
        Generate all three reports (summary, training, nutrition).
        
        Args:
            questionnaire_data: Parsed questionnaire data
            llm_content: Generated content from LLM
            chart_path: Path to radar chart image
            
        Returns:
            Dictionary with paths to generated reports
        """
        generated_reports = {}
        
        try:
            # Generate Summary Report
            summary_path = self.generate_summary_report(
                questionnaire_data, 
                llm_content.get("summary", ""), 
                chart_path
            )
            generated_reports["summary"] = summary_path
            
            # Generate Training Plan
            training_path = self.generate_training_plan(
                questionnaire_data,
                llm_content.get("training", "")
            )
            generated_reports["training"] = training_path
            
            # Generate Nutrition Plan
            nutrition_path = self.generate_nutrition_plan(
                questionnaire_data,
                llm_content.get("nutrition", "")
            )
            generated_reports["nutrition"] = nutrition_path
            
            self.logger.info("All reports generated successfully")
            return generated_reports
            
        except Exception as e:
            self.logger.error(f"Error generating reports: {str(e)}")
            raise
    
    def generate_summary_report(self, 
                               questionnaire_data: Dict[str, Any],
                               summary_content: str,
                               chart_path: Optional[str] = None) -> str:
        """
        Generate the summary report with radar chart.
        
        Args:
            questionnaire_data: Parsed questionnaire data
            summary_content: AI-generated summary content
            chart_path: Path to radar chart image
            
        Returns:
            Path to generated summary report
        """
        try:
            # Create document or load template
            doc = self._load_or_create_template("summary_template.docx")
            
            # Add header
            self._add_header(doc, "HEALTH & FITNESS SUMMARY REPORT")
            
            # Add personal information section
            self._add_personal_info_section(doc, questionnaire_data.get("personal_info", {}))
            
            # Add radar chart if available
            if chart_path and os.path.exists(chart_path):
                self._add_chart_section(doc, chart_path, "Health Assessment Overview")
            
            # Add scale responses summary
            self._add_scale_responses_section(doc, questionnaire_data.get("scale_responses", {}))
            
            # Add AI-generated summary
            self._add_content_section(doc, "Assessment Summary", summary_content)
            
            # Add goals and recommendations
            self._add_goals_section(doc, questionnaire_data.get("fitness_goals", []))
            
            # Add footer
            self._add_footer(doc)
            
            # Save document
            output_path = os.path.join(self.output_dir, f"Summary_Report_{self._get_timestamp()}.docx")
            doc.save(output_path)
            
            self.logger.info(f"Summary report generated: {output_path}")
            return output_path
            
        except Exception as e:
            self.logger.error(f"Error generating summary report: {str(e)}")
            raise
    
    def generate_training_plan(self, 
                              questionnaire_data: Dict[str, Any],
                              training_content: str) -> str:
        """
        Generate the personalized training plan.
        
        Args:
            questionnaire_data: Parsed questionnaire data
            training_content: AI-generated training content
            
        Returns:
            Path to generated training plan
        """
        try:
            doc = self._load_or_create_template("training_template.docx")
            
            # Add header
            self._add_header(doc, "PERSONALIZED TRAINING PLAN")
            
            # Add client information
            self._add_personal_info_section(doc, questionnaire_data.get("personal_info", {}))
            
            # Add fitness assessment
            self._add_fitness_assessment_section(doc, questionnaire_data)
            
            # Add AI-generated training plan
            self._add_content_section(doc, "Your 4-Week Training Program", training_content)
            
            # Add equipment and preferences
            self._add_equipment_section(doc, questionnaire_data.get("checkboxes", {}))
            
            # Add progress tracking section
            self._add_progress_tracking_section(doc)
            
            # Add footer
            self._add_footer(doc)
            
            # Save document
            output_path = os.path.join(self.output_dir, f"Training_Plan_{self._get_timestamp()}.docx")
            doc.save(output_path)
            
            self.logger.info(f"Training plan generated: {output_path}")
            return output_path
            
        except Exception as e:
            self.logger.error(f"Error generating training plan: {str(e)}")
            raise
    
    def generate_nutrition_plan(self, 
                               questionnaire_data: Dict[str, Any],
                               nutrition_content: str) -> str:
        """
        Generate the personalized nutrition plan.
        
        Args:
            questionnaire_data: Parsed questionnaire data
            nutrition_content: AI-generated nutrition content
            
        Returns:
            Path to generated nutrition plan
        """
        try:
            doc = self._load_or_create_template("nutrition_template.docx")
            
            # Add header
            self._add_header(doc, "PERSONALIZED NUTRITION PLAN")
            
            # Add client information
            self._add_personal_info_section(doc, questionnaire_data.get("personal_info", {}))
            
            # Add dietary preferences and restrictions
            self._add_dietary_preferences_section(doc, questionnaire_data.get("dietary_preferences", {}))
            
            # Add AI-generated nutrition plan
            self._add_content_section(doc, "Your Personalized Nutrition Guide", nutrition_content)
            
            # Add health metrics
            self._add_health_metrics_section(doc, questionnaire_data.get("health_metrics", {}))
            
            # Add footer
            self._add_footer(doc)
            
            # Save document
            output_path = os.path.join(self.output_dir, f"Nutrition_Plan_{self._get_timestamp()}.docx")
            doc.save(output_path)
            
            self.logger.info(f"Nutrition plan generated: {output_path}")
            return output_path
            
        except Exception as e:
            self.logger.error(f"Error generating nutrition plan: {str(e)}")
            raise
    
    def _load_or_create_template(self, template_name: str) -> Document:
        """Load existing template or create new document."""
        template_path = os.path.join(self.templates_dir, template_name)
        
        if os.path.exists(template_path):
            try:
                doc = Document(template_path)
                self.logger.info(f"Loaded template: {template_name}")
                return doc
            except Exception as e:
                self.logger.warning(f"Could not load template {template_name}: {str(e)}")
        
        # Create new document if template doesn't exist or can't be loaded
        doc = Document()
        self._setup_document_styles(doc)
        return doc
    
    def _setup_document_styles(self, doc: Document):
        """Set up document styles for consistent formatting."""
        try:
            # Define custom styles
            styles = doc.styles
            
            # Heading 1 style
            if 'Custom Heading 1' not in [s.name for s in styles]:
                heading1_style = styles.add_style('Custom Heading 1', WD_STYLE_TYPE.PARAGRAPH)
                heading1_font = heading1_style.font
                heading1_font.name = 'Arial'
                heading1_font.size = Pt(18)
                heading1_font.bold = True
                heading1_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                heading1_style.paragraph_format.space_after = Pt(12)
            
            # Heading 2 style
            if 'Custom Heading 2' not in [s.name for s in styles]:
                heading2_style = styles.add_style('Custom Heading 2', WD_STYLE_TYPE.PARAGRAPH)
                heading2_font = heading2_style.font
                heading2_font.name = 'Arial'
                heading2_font.size = Pt(14)
                heading2_font.bold = True
                heading2_style.paragraph_format.space_after = Pt(6)
            
            # Body text style
            if 'Custom Body' not in [s.name for s in styles]:
                body_style = styles.add_style('Custom Body', WD_STYLE_TYPE.PARAGRAPH)
                body_font = body_style.font
                body_font.name = 'Arial'
                body_font.size = Pt(11)
                body_style.paragraph_format.space_after = Pt(6)
                
        except Exception as e:
            self.logger.warning(f"Could not set up custom styles: {str(e)}")
    
    def _add_header(self, doc: Document, title: str):
        """Add document header with title."""
        # Add title
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(title)
        title_run.font.name = 'Arial'
        title_run.font.size = Pt(20)
        title_run.bold = True
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date
        date_paragraph = doc.add_paragraph()
        date_run = date_paragraph.add_run(f"Generated on: {datetime.now().strftime('%B %d, %Y')}")
        date_run.font.name = 'Arial'
        date_run.font.size = Pt(12)
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add separator line
        doc.add_paragraph("_" * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()  # Empty line
    
    def _add_personal_info_section(self, doc: Document, personal_info: Dict[str, str]):
        """Add personal information section."""
        if not personal_info:
            return
        
        # Section header
        heading = doc.add_paragraph()
        heading_run = heading.add_run("PERSONAL INFORMATION")
        heading_run.font.name = 'Arial'
        heading_run.font.size = Pt(14)
        heading_run.bold = True
        
        # Create a simple info layout
        info_items = [
            ("Name", personal_info.get("name", "Not provided")),
            ("Age", personal_info.get("age", "Not provided")),
            ("Gender", personal_info.get("gender", "Not provided")),
            ("Height", personal_info.get("height", "Not provided")),
            ("Weight", personal_info.get("weight", "Not provided"))
        ]
        
        for label, value in info_items:
            if value and value != "Not provided":
                para = doc.add_paragraph()
                para.add_run(f"{label}: ").bold = True
                para.add_run(value)
        
        doc.add_paragraph()  # Empty line
    
    def _add_chart_section(self, doc: Document, chart_path: str, section_title: str):
        """Add chart image to document."""
        try:
            # Section header
            heading = doc.add_paragraph()
            heading_run = heading.add_run(section_title.upper())
            heading_run.font.name = 'Arial'
            heading_run.font.size = Pt(14)
            heading_run.bold = True
            
            # Add chart image
            chart_paragraph = doc.add_paragraph()
            chart_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_picture(chart_path, width=Inches(6))
            
            doc.add_paragraph()  # Empty line
            
        except Exception as e:
            self.logger.warning(f"Could not add chart to document: {str(e)}")
            # Add placeholder text instead
            doc.add_paragraph("Chart could not be displayed.")
    
    def _add_scale_responses_section(self, doc: Document, scale_responses: Dict[str, int]):
        """Add scale responses summary."""
        if not scale_responses:
            return
        
        # Section header
        heading = doc.add_paragraph()
        heading_run = heading.add_run("ASSESSMENT SCORES (1-10 SCALE)")
        heading_run.font.name = 'Arial'
        heading_run.font.size = Pt(14)
        heading_run.bold = True
        
        # Add scores
        for category, score in scale_responses.items():
            para = doc.add_paragraph()
            category_formatted = category.replace('_', ' ').title()
            para.add_run(f"{category_formatted}: ").bold = True
            para.add_run(f"{score}/10")
        
        doc.add_paragraph()  # Empty line
    
    def _add_content_section(self, doc: Document, section_title: str, content: str):
        """Add AI-generated content section."""
        # Section header
        heading = doc.add_paragraph()
        heading_run = heading.add_run(section_title.upper())
        heading_run.font.name = 'Arial'
        heading_run.font.size = Pt(14)
        heading_run.bold = True
        
        # Add content
        if content:
            # Split content into paragraphs
            paragraphs = content.split('\n\n')
            for paragraph_text in paragraphs:
                if paragraph_text.strip():
                    para = doc.add_paragraph(paragraph_text.strip())
                    para.style.font.name = 'Arial'
                    para.style.font.size = Pt(11)
        else:
            doc.add_paragraph("Content will be generated based on your questionnaire responses.")
        
        doc.add_paragraph()  # Empty line
    
    def _add_goals_section(self, doc: Document, fitness_goals: list):
        """Add fitness goals section."""
        if not fitness_goals:
            return
        
        # Section header
        heading = doc.add_paragraph()
        heading_run = heading.add_run("YOUR FITNESS GOALS")
        heading_run.font.name = 'Arial'
        heading_run.font.size = Pt(14)
        heading_run.bold = True
        
        # Add goals as bullet points
        for goal in fitness_goals:
            para = doc.add_paragraph(goal.title(), style='List Bullet')
        
        doc.add_paragraph()  # Empty line
    
    def _add_fitness_assessment_section(self, doc: Document, questionnaire_data: Dict[str, Any]):
        """Add fitness assessment section for training plan."""
        scale_responses = questionnaire_data.get('scale_responses', {})
        if not scale_responses:
            return
        
        # Section header
        heading = doc.add_paragraph()
        heading_run = heading.add_run("FITNESS ASSESSMENT")
        heading_run.font.name = 'Arial'
        heading_run.font.size = Pt(14)
        heading_run.bold = True
        
        # Calculate average fitness level
        fitness_scores = [v for k, v in scale_responses.items() if 'fitness' in k.lower() or 'strength' in k.lower()]
        if fitness_scores:
            avg_fitness = sum(fitness_scores) / len(fitness_scores)
            level = "Beginner" if avg_fitness < 4 else "Intermediate" if avg_fitness < 7 else "Advanced"
            
            para = doc.add_paragraph()
            para.add_run("Assessed Fitness Level: ").bold = True
            para.add_run(f"{level} (Average Score: {avg_fitness:.1f}/10)")
        
        doc.add_paragraph()  # Empty line
    
    def _add_equipment_section(self, doc: Document, checkboxes: Dict[str, list]):
        """Add equipment and preferences section."""
        equipment = checkboxes.get('equipment_available', [])
        workout_times = checkboxes.get('workout_times', [])
        
        if equipment or workout_times:
            # Section header
            heading = doc.add_paragraph()
            heading_run = heading.add_run("EQUIPMENT & PREFERENCES")
            heading_run.font.name = 'Arial'
            heading_run.font.size = Pt(14)
            heading_run.bold = True
            
            if equipment:
                para = doc.add_paragraph()
                para.add_run("Available Equipment: ").bold = True
                para.add_run(", ".join(equipment))
            
            if workout_times:
                para = doc.add_paragraph()
                para.add_run("Preferred Workout Times: ").bold = True
                para.add_run(", ".join(workout_times))
            
            doc.add_paragraph()  # Empty line
    
    def _add_dietary_preferences_section(self, doc: Document, dietary_preferences: Dict[str, Any]):
        """Add dietary preferences section."""
        if not dietary_preferences:
            return
        
        # Section header
        heading = doc.add_paragraph()
        heading_run = heading.add_run("DIETARY PREFERENCES & RESTRICTIONS")
        heading_run.font.name = 'Arial'
        heading_run.font.size = Pt(14)
        heading_run.bold = True
        
        # Add restrictions
        restrictions = dietary_preferences.get('restrictions', [])
        if restrictions:
            para = doc.add_paragraph()
            para.add_run("Dietary Restrictions: ").bold = True
            para.add_run(", ".join(restrictions))
        
        # Add allergies
        allergies = dietary_preferences.get('allergies', [])
        if allergies:
            para = doc.add_paragraph()
            para.add_run("Allergies: ").bold = True
            para.add_run(", ".join(allergies))
        
        doc.add_paragraph()  # Empty line
    
    def _add_health_metrics_section(self, doc: Document, health_metrics: Dict[str, Any]):
        """Add health metrics section."""
        if not health_metrics:
            return
        
        # Section header
        heading = doc.add_paragraph()
        heading_run = heading.add_run("HEALTH METRICS")
        heading_run.font.name = 'Arial'
        heading_run.font.size = Pt(14)
        heading_run.bold = True
        
        # Add metrics
        for metric, value in health_metrics.items():
            if value:
                para = doc.add_paragraph()
                metric_formatted = metric.replace('_', ' ').title()
                para.add_run(f"{metric_formatted}: ").bold = True
                para.add_run(str(value))
        
        doc.add_paragraph()  # Empty line
    
    def _add_progress_tracking_section(self, doc: Document):
        """Add progress tracking section."""
        # Section header
        heading = doc.add_paragraph()
        heading_run = heading.add_run("PROGRESS TRACKING")
        heading_run.font.name = 'Arial'
        heading_run.font.size = Pt(14)
        heading_run.bold = True
        
        # Add tracking guidelines
        tracking_text = """
        Track your progress weekly by recording:
        • Workout completion and difficulty level
        • Weight lifted and repetitions achieved
        • Energy levels before and after workouts
        • Any modifications needed
        • Weekly body measurements (optional)
        
        Adjust your plan if you consistently find workouts too easy or too difficult.
        """
        
        doc.add_paragraph(tracking_text.strip())
        doc.add_paragraph()  # Empty line
    
    def _add_footer(self, doc: Document):
        """Add document footer."""
        # Add separator line
        doc.add_paragraph("_" * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add footer text
        footer_paragraph = doc.add_paragraph()
        footer_run = footer_paragraph.add_run(
            "This personalized plan was generated based on your questionnaire responses. "
            "Please consult with healthcare professionals before starting any new fitness or nutrition program."
        )
        footer_run.font.name = 'Arial'
        footer_run.font.size = Pt(10)
        footer_run.italic = True
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _get_timestamp(self) -> str:
        """Get timestamp for file naming."""
        return datetime.now().strftime("%Y%m%d_%H%M%S")
    
    def create_sample_templates(self):
        """Create sample Word templates for testing."""
        try:
            # Create basic templates
            template_names = ["summary_template.docx", "training_template.docx", "nutrition_template.docx"]
            
            for template_name in template_names:
                template_path = os.path.join(self.templates_dir, template_name)
                if not os.path.exists(template_path):
                    doc = Document()
                    self._setup_document_styles(doc)
                    
                    # Add placeholder content
                    doc.add_paragraph(f"Template: {template_name}")
                    doc.add_paragraph("This is a placeholder template. Content will be replaced during report generation.")
                    
                    doc.save(template_path)
                    self.logger.info(f"Created sample template: {template_name}")
            
        except Exception as e:
            self.logger.error(f"Error creating sample templates: {str(e)}")


if __name__ == "__main__":
    # Test the report generator
    logging.basicConfig(level=logging.INFO)
    
    print("Testing Report Generator...")
    print("=" * 50)
    
    # Create test directories
    os.makedirs("test_templates", exist_ok=True)
    os.makedirs("test_output", exist_ok=True)
    
    report_gen = ReportGenerator("test_templates", "test_output")
    report_gen.create_sample_templates()
    
    # Test data
    test_questionnaire_data = {
        "personal_info": {
            "name": "John Doe",
            "age": "30",
            "gender": "Male",
            "height": "6'0\"",
            "weight": "180 lbs"
        },
        "scale_responses": {
            "fitness_level": 6,
            "energy_level": 7,
            "motivation": 8
        },
        "fitness_goals": ["lose weight", "gain muscle"],
        "dietary_preferences": {
            "restrictions": ["vegetarian"],
            "allergies": ["nuts"]
        },
        "health_metrics": {
            "bmi": "24.5",
            "activity_level": "moderate"
        },
        "checkboxes": {
            "equipment_available": ["dumbbells", "yoga mat"],
            "workout_times": ["morning", "evening"]
        }
    }
    
    test_llm_content = {
        "summary": "Based on your assessment, you show good motivation and moderate fitness levels...",
        "training": "Your 4-week training program focuses on progressive overload...",
        "nutrition": "Your personalized nutrition plan emphasizes plant-based proteins..."
    }
    
    # Generate test reports
    try:
        reports = report_gen.generate_all_reports(test_questionnaire_data, test_llm_content)
        
        print("✓ Generated reports:")
        for report_type, path in reports.items():
            print(f"  - {report_type.title()}: {path}")
        
        print("\n✓ Report generation test completed successfully!")
        
    except Exception as e:
        print(f"✗ Error during testing: {e}")
